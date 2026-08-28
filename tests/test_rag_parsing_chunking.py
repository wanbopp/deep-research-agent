"""三类文档解析、Registry 与确定性分块的聚焦门禁."""

from hashlib import sha256
from io import BytesIO
from uuid import UUID

from docx import Document as DocxDocument
from pypdf import PdfWriter
from pypdf.generic import DecodedStreamObject, DictionaryObject, NameObject
import pytest

from app.rag.chunker import TokenAwareChunker
from app.rag.parsers import (
    DocumentParseError,
    DocxParser,
    MarkdownParser,
    ParseErrorCode,
    ParseRequest,
    ParsedBlockKind,
    ParserRegistry,
    PdfParser,
)


def _request(*, filename: str, content_type: str, content: bytes) -> ParseRequest:
    """用真实内容摘要构造固定解析请求，避免测试绕过生产契约."""
    return ParseRequest(
        filename=filename,
        content_type=content_type,
        content_sha256=sha256(content).hexdigest(),
        content=content,
    )


def _docx_bytes() -> bytes:
    """在内存中生成包含标题、中文正文和表格的真实 DOCX 包."""
    document = DocxDocument()
    document.add_heading("安装", level=1)
    document.add_paragraph("Windows 环境使用项目内部虚拟环境。")
    table = document.add_table(rows=1, cols=2)
    table.cell(0, 0).text = "命令"
    table.cell(0, 1).text = "说明"
    output = BytesIO()
    document.save(output)
    return output.getvalue()


def _pdf_bytes(*, with_text: bool) -> bytes:
    """生成真实 PDF；无文本版本用于证明扫描件不会被误判为空成功."""
    writer = PdfWriter()
    page = writer.add_blank_page(width=300, height=300)
    if with_text:
        font = DictionaryObject(
            {
                NameObject("/Type"): NameObject("/Font"),
                NameObject("/Subtype"): NameObject("/Type1"),
                NameObject("/BaseFont"): NameObject("/Helvetica"),
            }
        )
        font_ref = writer._add_object(font)  # noqa: SLF001 - 测试生成固定 PDF 对象。
        page[NameObject("/Resources")] = DictionaryObject(
            {NameObject("/Font"): DictionaryObject({NameObject("/F1"): font_ref})}
        )
        stream = DecodedStreamObject()
        stream.set_data(b"BT /F1 12 Tf 40 250 Td (DeepResearch PDF fact) Tj ET")
        page[NameObject("/Contents")] = writer._add_object(stream)  # noqa: SLF001
    output = BytesIO()
    writer.write(output)
    return output.getvalue()


@pytest.mark.anyio
async def test_real_parsers_registry_and_stable_failures() -> None:
    """Registry 应真实解析 Markdown/DOCX/PDF，并稳定识别 OCR 与未知格式."""
    registry = ParserRegistry((MarkdownParser(), DocxParser(), PdfParser()))
    markdown = b"# Guide\n\nParagraph\n\n```python\nprint('ok')\n```\n\n| A | B |"
    markdown_result = await registry.parse(
        _request(filename="guide.md", content_type="text/markdown", content=markdown)
    )
    docx_result = await registry.parse(
        _request(
            filename="guide.docx",
            content_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            content=_docx_bytes(),
        )
    )
    pdf_result = await registry.parse(
        _request(filename="guide.pdf", content_type="application/pdf", content=_pdf_bytes(with_text=True))
    )

    assert {block.kind for block in markdown_result.blocks} >= {
        ParsedBlockKind.HEADING,
        ParsedBlockKind.CODE,
        ParsedBlockKind.TABLE,
    }
    assert docx_result.blocks[1].location.section_path == ("安装",)
    assert pdf_result.blocks[0].location.page_number == 1

    with pytest.raises(DocumentParseError) as ocr_error:
        await registry.parse(
            _request(filename="scan.pdf", content_type="application/pdf", content=_pdf_bytes(with_text=False))
        )
    assert ocr_error.value.code is ParseErrorCode.OCR_REQUIRED

    with pytest.raises(DocumentParseError) as unsupported_error:
        registry.resolve("application/octet-stream")
    assert unsupported_error.value.code is ParseErrorCode.UNSUPPORTED

    with pytest.raises(ValueError, match="Duplicate parser"):
        ParserRegistry((MarkdownParser(), MarkdownParser()))


@pytest.mark.anyio
async def test_token_chunking_is_deterministic_bounded_and_keeps_provenance() -> None:
    """同一输入与配置必须得到相同 ID；中文长文本也必须受 token 预算约束."""
    content = ("# 中文章节\n\n" + "没有空格的中文事实" * 80).encode()
    request = _request(filename="facts.md", content_type="text/markdown", content=content)
    parsed = await MarkdownParser().parse(request)
    chunker = TokenAwareChunker(chunk_size=40, chunk_overlap=8)

    document_id = UUID("11111111-1111-4111-8111-111111111111")
    first = chunker.split(document_id=document_id, content_sha256=request.content_sha256, document=parsed)
    second = chunker.split(document_id=document_id, content_sha256=request.content_sha256, document=parsed)

    assert first == second
    assert len(first) > 1
    assert all(chunk.token_count <= 40 for chunk in first)
    assert [chunk.ordinal for chunk in first] == list(range(len(first)))
    assert all(chunk.sources for chunk in first)
    assert len({chunk.id for chunk in first}) == len(first)

    with pytest.raises(ValueError, match="chunk_overlap"):
        TokenAwareChunker(chunk_size=20, chunk_overlap=20)
