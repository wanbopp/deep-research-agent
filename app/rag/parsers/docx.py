"""基于 python-docx 的 DOCX parser adapter."""

import asyncio
from io import BytesIO
from zipfile import BadZipFile

from docx import Document as DocxDocument
from docx.opc.exceptions import PackageNotFoundError

from app.rag.parsers.contracts import (
    DocumentParseError,
    ParseErrorCode,
    ParseRequest,
    ParsedBlock,
    ParsedBlockKind,
    ParsedDocument,
    SourceLocation,
)
from app.rag.parsers.normalization import normalize_text


class DocxParser:
    """提取 DOCX 段落、标题层级和表格，不执行宏或外部链接."""

    name = "docx"
    version = "1"
    supported_content_types = frozenset({"application/vnd.openxmlformats-officedocument.wordprocessingml.document"})

    async def parse(self, request: ParseRequest) -> ParsedDocument:
        """把 ZIP/XML 解包和遍历放到线程，保护异步 worker 的事件循环."""
        return await asyncio.to_thread(self._parse_sync, request)

    def _parse_sync(self, request: ParseRequest) -> ParsedDocument:
        try:
            document = DocxDocument(BytesIO(request.content))
        except (BadZipFile, PackageNotFoundError, KeyError, ValueError):
            raise DocumentParseError(ParseErrorCode.CORRUPT) from None

        blocks: list[ParsedBlock] = []
        sections: dict[int, str] = {}
        offset = 0
        for paragraph in document.paragraphs:
            text = normalize_text(paragraph.text)
            if not text:
                continue
            kind = ParsedBlockKind.PARAGRAPH
            style_name = (paragraph.style.name or "") if paragraph.style is not None else ""
            if style_name.startswith("Heading "):
                try:
                    level = int(style_name.removeprefix("Heading "))
                except ValueError:
                    level = 1
                sections = {key: value for key, value in sections.items() if key < level}
                sections[level] = text
                kind = ParsedBlockKind.HEADING
            blocks.append(self._block(blocks, kind, text, sections, offset))
            offset += len(text) + 1

        for table in document.tables:
            rows = ["\t".join(normalize_text(cell.text) for cell in row.cells) for row in table.rows]
            text = normalize_text("\n".join(rows))
            if text:
                blocks.append(self._block(blocks, ParsedBlockKind.TABLE, text, sections, offset))
                offset += len(text) + 1

        if not blocks:
            raise DocumentParseError(ParseErrorCode.EMPTY)
        return ParsedDocument(
            parser_name=self.name,
            parser_version=self.version,
            content_type=request.content_type,
            blocks=tuple(blocks),
        )

    @staticmethod
    def _block(
        blocks: list[ParsedBlock],
        kind: ParsedBlockKind,
        text: str,
        sections: dict[int, str],
        offset: int,
    ) -> ParsedBlock:
        return ParsedBlock(
            ordinal=len(blocks),
            kind=kind,
            text=text,
            location=SourceLocation(
                section_path=tuple(sections[key] for key in sorted(sections)),
                source_start=offset,
                source_end=offset + len(text),
            ),
        )


__all__ = ["DocxParser"]
